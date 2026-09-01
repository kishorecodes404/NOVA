"""
Document Generation Agent
=========================

Generates real, downloadable .docx business documents assembled from
data that already lives inside NOVA (PO/leave/expense stores, the
calendar, indexed documents) - never from anything the LLM invents.

Design mirrors the rest of the app's "read agents":

    detect intent -> pull real records from rag.py -> build a
    grounded evidence block for the answer LLM -> separately, write
    an actual .docx file to disk so the UI can offer it as a
    download.

The evidence block (`context`) is fed to the SAME grounding-prompt
machinery `build_routed_prompt()` already uses for RECOMMEND/PLAN, so
the chat reply describing the document is held to the same
no-invented-facts bar.

FACTS vs WORDING, and why the document genuinely needs the AI:
    Every fact that appears in a generated document (names, dates,
    amounts, statuses, IDs) always comes straight from the real
    records pulled from rag.py in Python - never from the LLM, and
    never anything it wasn't given. The prose that carries those
    facts - the letter body, the summary paragraph on a PO/expense/
    minutes document - is written by the LLM via the `llm_writer`
    callback (see `_ai_generate_document_content()`), and that
    output is mechanically verified to contain every required fact
    before it is trusted. There is deliberately no hard-coded
    "fill-in-the-blanks" paragraph behind it: if `llm_writer` is
    unavailable, raises, or drops/alters a required fact, that
    document's generation FAILS with a clear error rather than
    silently completing from a static template - see
    `DocumentWriterUnavailable` below. This means the .docx's wording
    genuinely depends on the model; removing it removes the ability
    to produce the document, not just some polish.

    Tables of line-item/ledger data (PO items, expense rows, the
    meeting agenda) stay Python-built, not AI-authored - a table of
    numbers pulled straight from a record is not "wording" an LLM
    should be paraphrasing, and doing so would only add hallucination
    risk for zero benefit.

Layout/visual design (title block, metadata table, section headings,
spacing, signature block - see the "PROFESSIONAL DOCUMENT LAYOUT"
section below) is a reusable Python template shared by every document
type. That's a presentation template, not a content one: it has no
document-specific sentences baked into it, so it doesn't conflict
with the "no hard-coded content template" rule above - swap in a new
`doc_type_label`/sections and it fits a future document type too.

Public entry point: generate_document_evidence(question)
    -> (context: str, sources: list[str], file_info: dict | None)

file_info, when a document was actually produced, looks like:
    {"path": "<abs path>.docx", "filename": "...", "doc_type": "..."}
"""

import os
import re
import uuid
from datetime import datetime, date, timedelta

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import rag


class DocumentWriterUnavailable(Exception):
    """
    Raised when a document's wording must be genuinely AI-authored
    (see the module docstring) but the AI writer model (`llm_writer`)
    was not supplied, could not be reached, or did not return
    content that verifiably contains every required fact.

    This is intentionally NOT recovered with a hard-coded fallback
    paragraph - generate_document_evidence()'s existing try/except
    catches it and surfaces a plain "couldn't generate that document
    right now" message instead, so the system can never produce "the
    same complete document" by quietly filling a static template when
    the model is unavailable.
    """
    pass

GENERATED_DOCS_DIR = os.path.join(
    os.path.dirname(os.path.abspath(__file__)), "generated_documents"
)

# Single-tenant convention: every record in the PO/leave/expense
# stores is keyed against the literal string "me" (see rag.py) - that
# key is fine for lookups, but it must never be printed verbatim in a
# document a real person is going to read or sign. Set the actual
# name NOVA should print in generated letters here, or via the
# NOVA_USER_DISPLAY_NAME env var. Falls back to a clearly-a-placeholder
# label (rather than silently printing "me") if neither is set, so a
# missing config is obvious in the output instead of looking like a bug.
USER_DISPLAY_NAME = os.environ.get("NOVA_USER_DISPLAY_NAME", "").strip() or "[Employee Name Not Configured]"

# Same single-tenant convention as USER_DISPLAY_NAME above, for the
# other pieces of employee identity a formal Leave Letter may need to
# show. Unlike USER_DISPLAY_NAME these have NO placeholder fallback -
# a signature block with no name would look broken, but a Leave
# Letter with no designation/organization line is a perfectly normal,
# complete letter. So when unset, each is simply omitted rather than
# printed as "[Not Configured]" - this is what satisfies "only when
# available in the system" without ever inventing a value.
USER_DESIGNATION = os.environ.get("NOVA_USER_DESIGNATION", "").strip()

# Kept for config compatibility / potential future document types,
# but not printed on the current plain-office-letter Leave Letter
# format below - its employee detail block is Name/Designation/
# Organization/Date only, per the requested layout.
USER_EMPLOYEE_ID = os.environ.get("NOVA_USER_EMPLOYEE_ID", "").strip()

# Same convention again, for the Leave Letter's "Company/Organization"
# detail line - omitted (never invented) when unset.
USER_ORGANIZATION = os.environ.get("NOVA_USER_ORGANIZATION", "").strip()


def _is_system_generated_reason(reason):
    """
    True for internal/system-authored reason strings (e.g. the fixed
    string _build_autonomous_leave_plan in app.py attaches to leave
    requests it files on the user's behalf) that describe HOW the
    request was filed, not WHY the person is taking leave. These must
    never be printed as if a human wrote them in a formal letter.
    """
    if not reason:
        return False
    return reason.strip().lower().startswith("requested via nova")


def _ensure_output_dir():
    os.makedirs(GENERATED_DOCS_DIR, exist_ok=True)


# =========================================================
# INTENT DETECTION
# =========================================================

# Checked by app.py's route_query() BEFORE the read-only DOCUMENT
# (RAG-search) signals, and before SEND_MAIL's "draft an email"
# signals (those are matched earlier still, so no clash there).
# Deliberately requires a generation VERB ("generate"/"create"/
# "draft"/"prepare"/"make") together with a document-shaped noun, so
# an ordinary question like "what's my PO status" never lands here.
_GENERATION_VERBS = (
    "generate", "create", "draft", "prepare", "make", "produce",
    "write up", "put together",
)

DOCUMENT_TYPE_SIGNALS = {
    "purchase_order": (
        "po document", "po copy", "purchase order document",
        "purchase order copy", "purchase order form",
    ),
    "leave_letter": (
        "leave letter", "leave certificate", "leave application letter",
        "leave application", "leave approval letter",
    ),
    "expense_report": (
        "expense report", "expense statement", "reimbursement report",
        "reimbursement summary", "expense summary",
    ),
    "meeting_minutes": (
        "meeting minutes", "minutes of meeting", "mom document", " mom ",
    ),
}


def looks_like_document_generation_request(question):
    """
    Fast, keyword-based check used by route_query(). Returns True
    only when the message both (a) asks to GENERATE something and
    (b) names one of the known document shapes above.
    """

    q = " " + question.lower().strip() + " "

    has_verb = any(verb in q for verb in _GENERATION_VERBS)
    if not has_verb:
        return False

    for signals in DOCUMENT_TYPE_SIGNALS.values():
        if any(signal in q for signal in signals):
            return True

    # Generic "generate a/the business document" / "generate a
    # document" - or a bare follow-up like "draft a copy" right after
    # a document was already being discussed (e.g. right after filing
    # a leave request) - still counts. Handled downstream by
    # _detect_document_type() falling back to conversation context,
    # or asking what kind if nothing recent gives it away, rather
    # than route_query() missing this class of request entirely.
    if any(
        phrase in q
        for phrase in (
            " a document ", " a doc ", " a business document ",
            " a copy ", " a draft copy ", " draft copy ", " copy of ",
        )
    ):
        return True

    return False


# Keywords that hint at which document type a vague follow-up ("draft
# a copy", "make one for me") is actually about, when the message
# itself names no explicit type. Checked against recent conversation
# text, most recent turn first - see _infer_doc_type_from_context().
_CONTEXT_TYPE_HINTS = {
    "leave_letter": ("leave", "sick leave", "annual leave", "casual leave"),
    "purchase_order": ("purchase order", " po ", "vendor"),
    "expense_report": ("expense", "reimbursement"),
    "meeting_minutes": ("meeting", "minutes", "mom"),
}


def _infer_doc_type_from_context(conversation_history):
    """
    Best-effort fallback for a vague follow-up that names no document
    type of its own (e.g. "draft a copy" right after NOVA just filed
    a sick leave request). Scans recent conversation turns, most
    recent first, for a document-shaped topic that was just being
    discussed, so the reply matches what the person was clearly just
    doing instead of forcing a "which kind?" round-trip - or worse,
    falling through to an unrelated agent (see looks_like_document_
    generation_request()'s docstring). Only used when the CURRENT
    message carries no explicit signal of its own; never overrides one
    that does.
    """
    if not conversation_history:
        return None

    lines = conversation_history.lower().splitlines()
    for line in reversed(lines):
        for doc_type, hints in _CONTEXT_TYPE_HINTS.items():
            if any(hint in line for hint in hints):
                return doc_type
    return None


def _detect_document_type(question, conversation_history=""):
    q = " " + question.lower().strip() + " "
    for doc_type, signals in DOCUMENT_TYPE_SIGNALS.items():
        if any(signal in q for signal in signals):
            return doc_type
    return _infer_doc_type_from_context(conversation_history)


# =========================================================
# HELPERS
# =========================================================

def _new_docx_path(doc_type):
    _ensure_output_dir()
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"{doc_type}_{stamp}_{uuid.uuid4().hex[:6]}.docx"
    return os.path.join(GENERATED_DOCS_DIR, filename), filename


# =========================================================
# PROFESSIONAL DOCUMENT LAYOUT
#
# A single, reusable visual template shared by every document type -
# letterhead-style title, a light metadata table, consistently
# styled section headings, justified body text, and a signature
# block. Purely presentational (fonts/colors/spacing): it contains
# no document-specific sentences, so it carries no "content" for the
# no-hard-coded-template rule to apply to - see the module docstring.
# Adding a new document type in future means calling these same
# helpers with new section labels, not writing new layout code.
# =========================================================

_ACCENT_COLOR = RGBColor(0x1F, 0x3A, 0x5F)      # deep navy
_ACCENT_COLOR_HEX = "1F3A5F"
_MUTED_COLOR = RGBColor(0x5A, 0x5A, 0x5A)
_LABEL_COLOR = RGBColor(0x44, 0x44, 0x44)


def _new_professional_document(doc_type_label, subtitle=None):
    """Start a new .docx with the shared letterhead-style layout:
    page margins, base typography, a title in the accent color, an
    optional subtitle, and a horizontal rule underneath. Returns the
    Document, ready for metadata/section content to be appended."""
    doc = Document()

    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    # Trimmed from 0.9in - this was the single biggest contributor to
    # the dead space above the title. 0.5in is still a normal,
    # professional top margin (most business-letter templates run
    # 0.5-0.75in) - it just stops wasting a near-full extra inch
    # before any content starts.
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.9)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)

    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(2)
    title_run = title_p.add_run(doc_type_label)
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = _ACCENT_COLOR

    if subtitle:
        sub_p = doc.add_paragraph()
        sub_p.paragraph_format.space_before = Pt(0)
        sub_p.paragraph_format.space_after = Pt(6)
        sub_run = sub_p.add_run(subtitle)
        sub_run.font.size = Pt(10.5)
        sub_run.font.color.rgb = _MUTED_COLOR
        sub_run.font.italic = True

    _add_horizontal_rule(doc)
    return doc


def _add_horizontal_rule(doc):
    """A clean bottom-border rule (never a table-as-a-line hack)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), _ACCENT_COLOR_HEX)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _add_metadata_block(doc, pairs):
    """Renders a clean, borderless label/value metadata table (Date,
    Reference No, Status, ...). `pairs` is a list of (label, value);
    entries with an empty value are skipped. Every value here comes
    straight from a real record, never invented."""
    visible = [(label, value) for label, value in pairs if value not in (None, "")]
    if not visible:
        return None

    table = doc.add_table(rows=0, cols=2)
    table.autofit = True
    for label, value in visible:
        row = table.add_row().cells
        # Word's default table cell margins add their own visible
        # padding above/below every row - trimmed here so a 3-4 row
        # metadata block doesn't read as taller than it needs to be.
        for cell in row:
            cell_pr = cell._tc.get_or_add_tcPr()
            margins = OxmlElement("w:tcMar")
            for side, value_twips in (("top", "20"), ("bottom", "20")):
                node = OxmlElement(f"w:{side}")
                node.set(qn("w:w"), value_twips)
                node.set(qn("w:type"), "dxa")
                margins.append(node)
            cell_pr.append(margins)
        label_run = row[0].paragraphs[0].add_run(str(label))
        label_run.bold = True
        label_run.font.size = Pt(10.5)
        label_run.font.color.rgb = _LABEL_COLOR
        value_run = row[1].paragraphs[0].add_run(str(value))
        value_run.font.size = Pt(10.5)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(4)
    return table


def _add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11.5)
    run.font.color.rgb = _ACCENT_COLOR
    return p


def _add_body_text(doc, text, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """Renders AI-authored (or otherwise free-form) prose as one or
    more paragraphs, splitting the passage on blank lines so a
    multi-paragraph AI response reads as real paragraphs rather than
    one dense block. `alignment` defaults to justified (the look used
    by the Purchase Order/Expense Report/Meeting Minutes documents);
    pass WD_ALIGN_PARAGRAPH.LEFT for a plain left-aligned letter
    style instead."""
    paragraphs = [chunk.strip() for chunk in text.split("\n\n") if chunk.strip()]
    for chunk in paragraphs or [text.strip()]:
        p = doc.add_paragraph(chunk)
        p.alignment = alignment
        p.paragraph_format.space_after = Pt(10)


def _add_signature_block(doc, name, closing="Regards,", extra_lines=None):
    """extra_lines: optional list of strings (e.g. designation,
    employee ID) printed one per line directly under the name, in the
    order given. Callers should only pass values that actually came
    from real config/records - each entry is simply skipped if it's
    empty, so a signature block never grows a blank line for a fact
    the system doesn't have."""
    lead_in = doc.add_paragraph(closing)
    lead_in.paragraph_format.space_before = Pt(10)
    lead_in.paragraph_format.space_after = Pt(28)  # room for a physical signature

    name_p = doc.add_paragraph()
    name_run = name_p.add_run(name)
    name_run.bold = True

    for line in extra_lines or []:
        if not line:
            continue
        extra_p = doc.add_paragraph()
        extra_p.paragraph_format.space_before = Pt(0)
        extra_run = extra_p.add_run(str(line))
        extra_run.font.size = Pt(10.5)
        extra_run.font.color.rgb = _MUTED_COLOR


def _match_vendor(question, po_history):
    """Best-effort match of a vendor name the user typed against
    real vendors in this user's PO history."""
    q = question.lower()
    for record in po_history:
        vendor = (record.get("vendor") or "").strip()
        if vendor and vendor.lower() in q:
            return record
    return None


def _match_leave_type(question):
    q = question.lower()
    for leave_type in ("annual", "sick", "casual"):
        if leave_type in q:
            return leave_type
    return None


def _format_letter_date(value):
    """Render a stored 'YYYY-MM-DD' leave date as 'September 02,
    2026' for the printed letter. Falls back to the raw stored value
    unchanged if it isn't in that shape (never invents a date)."""
    if not value:
        return ""
    try:
        return datetime.strptime(str(value), "%Y-%m-%d").strftime("%B %d, %Y")
    except ValueError:
        return str(value)


def _short_reference(record_id, prefix):
    """
    Formats a record's real (internally a raw UUID4) id into a
    compact, document-appropriate reference code, e.g. 'LV-0D79F746'.
    Uses only a leading substring of the actual retrieved id - it
    never invents a new identifier - so the printed reference still
    traces back to the real record; the full id remains available in
    evidence_lines/sources for audit purposes.
    """
    raw = str(record_id or "").replace("-", "")
    short = raw[:8].upper() if raw else "UNKNOWN"
    return f"{prefix}-{short}"


# =========================================================
# AI-GENERATED CONTENT (facts injected exactly, never invented)
#
# The one place in this module that talks to a model. Every builder
# below calls this - never a hard-coded f-string paragraph - to
# produce the actual wording that goes in the document. Reusable
# as-is for any future document type: pass a doc-type label, an
# instruction, and the real facts it must contain.
# =========================================================

# Standalone lines a model sometimes adds anyway despite being told
# the greeting/closing/signature are handled separately by the
# layout (observed in practice with local models at temp 0.4, which
# tend to default to "write a complete letter" habits). Matched only
# as an ENTIRE stripped line - never a substring inside a real
# sentence - so genuine AI-authored content is never touched, only
# accidental duplicate boilerplate is removed.
_BOILERPLATE_LINE_RE = re.compile(
    r"^(to whom it may concern|dear\s+[^,:\n]{0,60}|"
    r"sincerely|regards|best regards|warm regards|kind regards|"
    r"yours (?:faithfully|sincerely|truly)|"
    r"\[[^\]]{0,60}\])\s*[:,]?\s*$",
    re.IGNORECASE,
)
_SIGNOFF_KEYWORD_RE = re.compile(
    r"^(sincerely|regards|best regards|warm regards|kind regards|"
    r"yours (?:faithfully|sincerely|truly))\s*[:,]?\s*$",
    re.IGNORECASE,
)


def _strip_boilerplate(text):
    """
    Defense-in-depth for the instruction above: strips a leading
    greeting line and a trailing sign-off (+ the name/placeholder
    line right after it, e.g. "Sincerely,\\n[Your Name]" or
    "Regards,\\nKishore M S") if the model wrote its own despite
    being told not to. Only ever removes lines that are ENTIRELY a
    known greeting/closing/placeholder pattern, or a short name-like
    line immediately following a removed sign-off keyword - it never
    alters or removes a line that carries real sentence content.
    """
    lines = text.split("\n")
    kept = []
    skip_next_if_name = False
    for line in lines:
        stripped = line.strip()
        if skip_next_if_name:
            skip_next_if_name = False
            if stripped and len(stripped) <= 60 and not stripped.endswith((".", "!", "?")):
                continue  # drop the model's own signature name
        if _BOILERPLATE_LINE_RE.match(stripped):
            if _SIGNOFF_KEYWORD_RE.match(stripped):
                skip_next_if_name = True
            continue
        kept.append(line)

    cleaned = "\n".join(kept)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    return cleaned.strip()


def _ai_generate_document_content(llm_writer, doc_type_label, instruction,
                                   required_facts, max_chars=1600):
    """
    Ask an LLM to write the prose content for a section of a formal
    `doc_type_label`, using ONLY the facts listed in `required_facts`
    (a list of (label, value) pairs) - the model chooses the wording,
    structure, and tone, but the facts themselves are never left to
    it to get right.

    Returns the AI's text - after stripping any duplicate greeting/
    closing boilerplate it added despite instructions not to, see
    _strip_boilerplate() - only if EVERY required fact's value still
    literally appears (case-insensitive) somewhere in what remains.
    Otherwise returns None - callers must treat None as "this
    document's content could not be produced" (see
    DocumentWriterUnavailable), never substitute their own template
    sentence, so the document's wording keeps genuinely depending on
    the model rather than a static fallback.

    llm_writer: a callable(prompt: str) -> str supplied by the
    caller (app.py's _call_document_writer_model), so this module
    never has to import a model client of its own. If llm_writer is
    None, or it raises, or its output fails verification, this
    returns None.
    """
    if llm_writer is None:
        return None

    facts_block = "\n".join(f"- {label}: {value}" for label, value in required_facts)

    prompt = f"""You are a professional business-writing assistant drafting the body content of a formal {doc_type_label}.

TASK: {instruction}

You MUST include every one of these facts, stated exactly as given -
do not alter any number, date, name, or status word:
{facts_block}

Rules:
- Use ONLY the facts listed above. Do NOT invent, assume, or add any other name, date, amount, or detail.
- Write natural, professional prose - do not just restate the facts as a list.
- The document already has its own greeting/salutation and its own closing/signature, added separately by the template - your output is ONLY the content that goes between them.
- Do NOT write a greeting or salutation (e.g. "To Whom It May Concern," "Dear ...,") - it is already there.
- Do NOT write a closing, sign-off, or signature (e.g. "Sincerely," "Regards," a name, or a placeholder like "[Your Name]") - it is already there.
- Do NOT add a title or heading.
- Do NOT use markdown, bullet points, or headings.
- Begin directly with the first sentence of the content itself, and end after the last sentence of the content itself - nothing before or after it.
- If more than one paragraph is genuinely needed, separate paragraphs with a single blank line.
- Keep it concise and formal.
- Output ONLY that content, nothing else.

Content:"""

    try:
        raw = llm_writer(prompt)
    except Exception:
        return None

    if not raw:
        return None

    text = _strip_boilerplate(raw.strip())

    if not text or len(text) > max_chars:
        return None

    lowered = text.lower()
    for _label, value in required_facts:
        if str(value).strip().lower() not in lowered:
            return None

    return text


# =========================================================
# PURCHASE ORDER DOCUMENT
# =========================================================

def _build_purchase_order_docx(question, llm_writer=None):
    # The line-item table and the requester's own Justification text
    # stay Python-built/verbatim (see module docstring) - but the
    # cover summary paragraph is genuinely AI-authored from the real
    # PO record, verified to contain the facts that matter (ID,
    # vendor, department, status). The exact monetary total is
    # intentionally NOT a required fact here (it's already shown
    # precisely in the metadata block and the totals line below), so
    # a model's natural phrasing of a number never has to be
    # character-matched.
    po_history = rag.get_po_history("me", include_cancelled=True)

    if not po_history:
        return None, "No purchase order requests were found for you to generate a document from.", [], None

    record = _match_vendor(question, po_history) or po_history[0]

    po_id = record.get("id", "")
    vendor = record.get("vendor", "")
    department = record.get("department", "")
    status = record.get("status", "")
    requested_at = record.get("requested_at", "")
    total_amount = record.get("total_amount", 0)

    required_facts = [
        ("PO ID", po_id),
        ("vendor", vendor),
        ("department", department),
        ("status", status),
    ]
    instruction = (
        f"Write a short cover summary (1-2 sentences) for a formal "
        f"Purchase Order document, stating that PO {po_id} was raised "
        f"for vendor {vendor} on behalf of the {department} "
        f"department, and that it currently has status {status}. The "
        f"exact line items and total are itemized in the table that "
        f"follows, so do not restate specific amounts."
    )
    summary = _ai_generate_document_content(
        llm_writer, "Purchase Order", instruction, required_facts
    )
    if not summary:
        raise DocumentWriterUnavailable(
            "the AI writer model is currently unavailable, so this "
            "Purchase Order's summary could not be generated"
        )

    path, filename = _new_docx_path("purchase_order")
    doc = _new_professional_document(
        "Purchase Order", subtitle=f"Reference: {_short_reference(po_id, 'PO')}"
    )

    _add_metadata_block(doc, [
        ("Requester", record.get("requester", "")),
        ("Vendor", vendor),
        ("Vendor Email", record.get("vendor_email", "")),
        ("Department", department),
        ("Status", status),
        ("Requested At", requested_at),
        ("Due Date", record.get("due_date", "")),
    ])

    _add_body_text(doc, summary)

    _add_section_heading(doc, "Line Items")
    items = record.get("items") or []
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = (
        "Item", "Quantity", "Unit Price (₹)", "Line Total (₹)"
    )
    for item in items:
        row = table.add_row().cells
        row[0].text = str(item.get("name", ""))
        row[1].text = rag.format_po_quantity(item.get("quantity", 0))
        row[2].text = f"{item.get('unit_price', 0):,.2f}"
        row[3].text = f"{item.get('line_total', 0):,.2f}"

    total_p = doc.add_paragraph()
    total_p.paragraph_format.space_before = Pt(8)
    total_run = total_p.add_run(f"Total Amount: ₹{total_amount:,.2f}")
    total_run.bold = True
    total_run.font.size = Pt(12)
    total_run.font.color.rgb = _ACCENT_COLOR

    if record.get("justification"):
        _add_section_heading(doc, "Justification")
        doc.add_paragraph(record["justification"])

    _add_signature_block(doc, record.get("requester", USER_DISPLAY_NAME))

    doc.save(path)

    evidence_lines = [
        f"PO_ID: {po_id}",
        f"VENDOR: {vendor}",
        f"DEPARTMENT: {department}",
        f"STATUS: {status}",
        f"TOTAL_AMOUNT: ₹{total_amount:,.2f}",
        f"REQUESTED_AT: {requested_at}",
        "SUMMARY_SOURCE: AI-written (fact-verified)",
    ]
    sources = [f"PO request {po_id} - {vendor} (₹{total_amount:,.2f})"]

    confirmation = (
        f"Your Purchase Order document is ready. It covers the PO for "
        f"**{vendor}** (₹{total_amount:,.2f}, "
        f"{department} department), currently "
        f"**{status}**."
    )

    return (path, filename), evidence_lines, sources, confirmation


# =========================================================
# LEAVE LETTER
# =========================================================

def _build_leave_letter_docx(question, llm_writer=None):
    leave_type = _match_leave_type(question)
    history = rag.get_leave_history("me", include_cancelled=True)

    if leave_type:
        history = [r for r in history if r.get("leave_type") == leave_type] or history

    if not history:
        return None, "No leave requests were found for you to generate a letter from.", [], None

    record = history[0]

    start_fmt = _format_letter_date(record.get("start"))
    end_fmt = _format_letter_date(record.get("end"))
    single_day = record.get("start") == record.get("end")
    when_line = f"on {start_fmt}" if single_day else f"from {start_fmt} to {end_fmt}"

    days = record.get("days", "?")
    day_word = "day" if days == 1 else "days"
    status = record.get("status", "pending")
    leave_type_label = record.get("leave_type", "").title()

    # Only a real, employee-given reason belongs in a formal letter -
    # a system-authored bookkeeping string (e.g. from an autonomous
    # task run) describes how the request was FILED, not why the
    # person is taking leave.
    reason = record.get("reason")
    show_reason = bool(reason) and not _is_system_generated_reason(reason)

    # The facts below are exactly what must appear in the letter,
    # word-for-word - the AI writer chooses how to phrase the body,
    # but never what facts go in it. There is no deterministic
    # fallback paragraph: if the model isn't available, or drops/
    # alters any of these, generation fails outright (see
    # DocumentWriterUnavailable) rather than quietly completing the
    # letter from a template - see module docstring.
    required_facts = [
        ("employee name", USER_DISPLAY_NAME),
        ("leave type", leave_type_label),
        ("start date", start_fmt),
        ("status", status),
        ("number of working days", str(days)),
    ]
    if not single_day:
        required_facts.append(("end date", end_fmt))
    if show_reason:
        required_facts.append(("reason", reason))

    instruction = (
        f"Write the body of a formal Leave Letter (the recipient is "
        f"addressed separately as 'To Whom It May Concern' - do not "
        f"repeat that greeting). The body must confirm that "
        f"{USER_DISPLAY_NAME} has requested {leave_type_label} Leave "
        f"{when_line}, covering {days} working {day_word}, and state "
        f"that the request is currently {status}. Close with a short "
        f"line offering to provide further information or "
        f"documentation if required."
        + (f" Naturally work in the reason given: {reason}." if show_reason else "")
    )

    body = _ai_generate_document_content(
        llm_writer, "Leave Letter", instruction, required_facts
    )
    if not body:
        raise DocumentWriterUnavailable(
            "the AI writer model is currently unavailable, so this "
            "Leave Letter's content could not be generated"
        )

    path, filename = _new_docx_path("leave_letter")
    today_str = datetime.now().strftime("%B %d, %Y")
    dates_phrase = start_fmt if single_day else f"{start_fmt} to {end_fmt}"

    # -----------------------------------------------------------
    # Plain office-letter presentation (deliberately NOT the shared
    # _new_professional_document/_add_metadata_block/
    # _add_horizontal_rule layout used by the PO/Expense Report/
    # Meeting Minutes documents - this is a distinct, simpler style
    # requested specifically for the Leave Letter: no colored rule,
    # no table, no shading, nothing decorative. Every value placed on
    # the page below still comes straight from real data/config,
    # exactly as before - only how it's laid out has changed.
    # -----------------------------------------------------------
    doc = Document()

    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.2
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)

    # Title - centered, bold, dark navy. The only centered element in
    # the document; everything else below is left-aligned.
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(20)
    title_run = title_p.add_run("Leave Application for Office")
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = _ACCENT_COLOR

    # Employee detail block - plain left-aligned lines, no table, no
    # borders/shading. Designation/Organization are printed only when
    # actually configured (USER_DESIGNATION/USER_ORGANIZATION) - never
    # invented for a letter that doesn't have them.
    detail_lines = [USER_DISPLAY_NAME]
    if USER_DESIGNATION:
        detail_lines.append(USER_DESIGNATION)
    if USER_ORGANIZATION:
        detail_lines.append(USER_ORGANIZATION)
    detail_lines.append(today_str)

    for line in detail_lines:
        detail_p = doc.add_paragraph()
        detail_p.paragraph_format.space_after = Pt(2)
        detail_p.add_run(line)

    # Subject line - deterministic, Python-built (never AI-authored),
    # assembled from the same real dates already validated above.
    subject_line = f"Subject: Leave Application for {dates_phrase}"
    subject_p = doc.add_paragraph()
    subject_p.paragraph_format.space_before = Pt(16)
    subject_p.paragraph_format.space_after = Pt(16)
    subject_run = subject_p.add_run(subject_line)
    subject_run.bold = True

    # Salutation
    salutation_p = doc.add_paragraph("To Whom It May Concern,")
    salutation_p.paragraph_format.space_after = Pt(14)

    # AI-generated body - left-aligned (not justified) to match the
    # plain office-letter style; the facts it must contain were
    # verified back in _ai_generate_document_content() above.
    _add_body_text(doc, body, alignment=WD_ALIGN_PARAGRAPH.LEFT)

    _add_signature_block(doc, USER_DISPLAY_NAME)

    doc.save(path)

    evidence_lines = [
        f"LEAVE_ID: {record.get('id')}",
        f"USER: {record.get('user')}",
        f"LEAVE_TYPE: {record.get('leave_type')}",
        f"DATES: {record.get('start')} to {record.get('end')}",
        f"DAYS: {record.get('days')}",
        f"STATUS: {record.get('status')}",
        f"SUBJECT: {subject_line}",
    ]
    if USER_DESIGNATION:
        evidence_lines.append(f"DESIGNATION: {USER_DESIGNATION}")
    if USER_ORGANIZATION:
        evidence_lines.append(f"ORGANIZATION: {USER_ORGANIZATION}")
    evidence_lines.append("LETTER_BODY_SOURCE: AI-written (fact-verified)")
    sources = [
        f"Leave request {record.get('id')} - {record.get('leave_type')} "
        f"({record.get('start')} to {record.get('end')})"
    ]

    confirmation = (
        f"Your Leave Letter is ready. It covers **{leave_type_label}** "
        f"leave from **{start_fmt}** to **{end_fmt}** "
        f"({days} {day_word}), status **{status}**."
    )

    return (path, filename), evidence_lines, sources, confirmation


# =========================================================
# EXPENSE REPORT (aggregate, not a single record)
# =========================================================

def _resolve_report_window(question):
    """Very small date-window resolver: 'this month' (default),
    'last month', or 'last N days'."""

    q = question.lower()
    today = date.today()

    match = re.search(r"last (\d+) days", q)
    if match:
        days = int(match.group(1))
        return today - timedelta(days=days), today, f"last {days} days"

    if "last month" in q:
        first_of_this_month = today.replace(day=1)
        last_month_end = first_of_this_month - timedelta(days=1)
        last_month_start = last_month_end.replace(day=1)
        return last_month_start, last_month_end, "last month"

    # Default: current month to date.
    return today.replace(day=1), today, "this month"


def _build_expense_report_docx(question, llm_writer=None):
    # The claims table itself stays Python-built (it's a ledger of
    # numbers, not prose) - the cover summary above it is genuinely
    # AI-authored from the real, retrieved claims for this window.
    # The exact total is intentionally not a required fact here (it's
    # already shown precisely in the totals line and every row of the
    # table), so the model's natural phrasing of a number is never
    # character-matched.
    start_window, end_window, window_label = _resolve_report_window(question)
    history = rag.get_expense_history("me", include_cancelled=False)

    in_window = [
        record for record in history
        if record.get("date_incurred")
        and start_window.isoformat() <= record["date_incurred"] <= end_window.isoformat()
    ]

    path, filename = _new_docx_path("expense_report")

    if not in_window:
        doc = _new_professional_document(
            "Expense Report", subtitle=f"Period: {window_label}"
        )
        _add_metadata_block(doc, [
            ("Employee", USER_DISPLAY_NAME),
            ("Period", f"{start_window} to {end_window} ({window_label})"),
        ])
        doc.add_paragraph("No expense claims were recorded in this period.")
        doc.save(path)
        evidence_lines = [f"PERIOD: {window_label}", "EXPENSE_COUNT: 0"]
        sources = [f"Expense history for me, {window_label} - no records found"]
        confirmation = (
            f"Your Expense Report is ready for **{window_label}**, though no "
            f"expense claims were found in that period."
        )
        return (path, filename), evidence_lines, sources, confirmation

    total = sum(r.get("amount", 0) or 0 for r in in_window)
    categories = sorted({
        str(r.get("category", "")).replace("_", " ").title()
        for r in in_window if r.get("category")
    })

    required_facts = [
        ("period", window_label),
        ("number of claims", str(len(in_window))),
    ]
    instruction = (
        f"Write a short cover summary (1-2 sentences) for a formal "
        f"Expense Report, stating that it covers {window_label} and "
        f"summarizes {len(in_window)} submitted expense claim(s) "
        f"across categories: {', '.join(categories)}. The itemized "
        f"claims and total are in the table that follows, so do not "
        f"restate specific amounts."
    )
    summary = _ai_generate_document_content(
        llm_writer, "Expense Report", instruction, required_facts
    )
    if not summary:
        raise DocumentWriterUnavailable(
            "the AI writer model is currently unavailable, so this "
            "Expense Report's summary could not be generated"
        )

    doc = _new_professional_document(
        "Expense Report", subtitle=f"Period: {window_label}"
    )
    _add_metadata_block(doc, [
        ("Employee", USER_DISPLAY_NAME),
        ("Period", f"{start_window} to {end_window} ({window_label})"),
        ("Claims Submitted", str(len(in_window))),
    ])

    _add_body_text(doc, summary)

    _add_section_heading(doc, "Itemized Claims")
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text, hdr[4].text = (
        "Date", "Category", "Vendor", "Amount (₹)", "Status"
    )
    for record in in_window:
        row = table.add_row().cells
        row[0].text = str(record.get("date_incurred", ""))
        row[1].text = str(record.get("category", "")).replace("_", " ").title()
        row[2].text = str(record.get("vendor", "") or "-")
        row[3].text = f"{record.get('amount', 0):,.2f}"
        row[4].text = str(record.get("status", ""))

    total_p = doc.add_paragraph()
    total_p.paragraph_format.space_before = Pt(8)
    total_run = total_p.add_run(f"Total Claimed: ₹{total:,.2f}")
    total_run.bold = True
    total_run.font.size = Pt(12)
    total_run.font.color.rgb = _ACCENT_COLOR

    _add_signature_block(doc, USER_DISPLAY_NAME)

    doc.save(path)

    evidence_lines = [
        f"PERIOD: {window_label}", f"EXPENSE_COUNT: {len(in_window)}",
        f"TOTAL_CLAIMED: ₹{total:,.2f}", "SUMMARY_SOURCE: AI-written (fact-verified)",
    ]
    evidence_lines += [
        f"EXPENSE: {r.get('date_incurred')} | {r.get('category')} | "
        f"₹{r.get('amount', 0):,.2f} | {r.get('status')}"
        for r in in_window
    ]
    sources = [
        f"Expense claim {r.get('id')} - {r.get('category')} ₹{r.get('amount', 0):,.2f}"
        for r in in_window
    ]

    confirmation = (
        f"Your Expense Report is ready for **{window_label}**: "
        f"{len(in_window)} claim(s) totaling **₹{total:,.2f}**."
    )

    return (path, filename), evidence_lines, sources, confirmation


# =========================================================
# MEETING MINUTES
# =========================================================

def _build_meeting_minutes_docx(question, llm_writer=None):
    # The agenda stays a straight bullet list of real calendar
    # events (data, not prose), and the Notes section is
    # deliberately left blank for a human to fill in after the
    # meeting actually happens - there's no narrative an AI could
    # faithfully pre-write for a discussion that hasn't occurred yet.
    # The AI's role is limited to a short, factual opening summary of
    # the day's agenda.
    target_date = date.today()
    events = rag.get_events_on_date(target_date, user="me")
    date_label = target_date.strftime("%A, %B %d, %Y")

    path, filename = _new_docx_path("meeting_minutes")

    if not events:
        doc = _new_professional_document("Meeting Minutes")
        _add_metadata_block(doc, [("Date", date_label)])
        doc.add_paragraph("No meetings were found on the calendar for this date.")
        doc.save(path)
        evidence_lines = [f"DATE: {target_date}", "MEETING_COUNT: 0"]
        sources = [f"Calendar for {target_date} - no events found"]
        confirmation = (
            f"Your Meeting Minutes document is ready for "
            f"**{target_date.strftime('%B %d, %Y')}**, though no meetings "
            f"were found on the calendar for that date."
        )
        return (path, filename), evidence_lines, sources, confirmation

    required_facts = [("date", date_label), ("number of meetings", str(len(events)))]
    instruction = (
        f"Write a short opening summary (1-2 sentences) for a formal "
        f"Meeting Minutes document covering {date_label}, noting that "
        f"there are {len(events)} agenda item(s) for the day. The "
        f"agenda itself is listed separately below, so do not restate "
        f"each event by name."
    )
    summary = _ai_generate_document_content(
        llm_writer, "Meeting Minutes", instruction, required_facts
    )
    if not summary:
        raise DocumentWriterUnavailable(
            "the AI writer model is currently unavailable, so this "
            "Meeting Minutes document's summary could not be generated"
        )

    doc = _new_professional_document("Meeting Minutes")
    _add_metadata_block(doc, [
        ("Date", date_label),
        ("Agenda Items", str(len(events))),
    ])

    _add_body_text(doc, summary)

    _add_section_heading(doc, "Agenda / Events")
    for label in events:
        doc.add_paragraph(label, style="List Bullet")

    _add_section_heading(doc, "Notes")
    doc.add_paragraph(
        "(Add discussion notes and action items for each agenda item above.)"
    )

    doc.save(path)

    evidence_lines = [
        f"DATE: {target_date}", f"MEETING_COUNT: {len(events)}",
        "SUMMARY_SOURCE: AI-written (fact-verified)",
    ]
    evidence_lines += [f"EVENT: {label}" for label in events]
    sources = [f"Calendar event - {label}" for label in events]

    confirmation = (
        f"Your Meeting Minutes document is ready for "
        f"**{target_date.strftime('%B %d, %Y')}**, covering "
        f"**{len(events)}** meeting(s) from your calendar."
    )

    return (path, filename), evidence_lines, sources, confirmation


_BUILDERS = {
    "purchase_order": _build_purchase_order_docx,
    "leave_letter": _build_leave_letter_docx,
    "expense_report": _build_expense_report_docx,
    "meeting_minutes": _build_meeting_minutes_docx,
}

_DOC_TYPE_LABEL = {
    "purchase_order": "Purchase Order",
    "leave_letter": "Leave Letter",
    "expense_report": "Expense Report",
    "meeting_minutes": "Meeting Minutes",
}


# =========================================================
# PUBLIC ENTRY POINT
# =========================================================

def generate_document_evidence(question, conversation_history="", llm_writer=None):
    """
    Returns (context, sources, file_info, confirmation).

    context: grounded evidence text - kept for the "Sources"-style
        audit trail and for any caller that still wants to run it
        through an LLM, but NOT required to produce an answer.
    sources: list[str] for the UI's "Sources" expander.
    file_info: {"path", "filename", "doc_type"} if a .docx was
        actually written, else None.
    confirmation: a ready-to-display, fully deterministic sentence
        confirming what was generated - built directly from the real
        record in Python, never by an LLM. app.py uses this AS THE
        ANSWER for the GENERATE_DOCUMENT route instead of routing
        through Ollama/Groq/Gemini, so a local-model outage (e.g. the
        Ollama server erroring on /api/generate) can never strand an
        already-generated file behind a failed chat reply. None only
        when no document type could be identified at all.

    conversation_history: recent chat text, used only as a fallback
        to infer the document type for a vague follow-up like "draft
        a copy" that names no type of its own - see
        _infer_doc_type_from_context(). Never overrides an explicit
        type named in `question` itself.

    llm_writer: callable(prompt: str) -> str, used to write the
        actual PROSE of the document (the Leave Letter's body, and
        the cover summary on PO/Expense Report/Meeting Minutes) - see
        _ai_generate_document_content(). Every fact that ends up in
        the document (names/dates/amounts/status) always comes
        straight from the real record in Python, never from this
        model; its output is verified to contain every required fact
        before it's trusted. Unlike a template-filling parameter,
        this is not optional in practice: pass None (or a writer that
        fails) and any document type whose wording depends on it will
        raise DocumentWriterUnavailable, caught below and surfaced as
        a plain failure message - there is no deterministic template
        it falls back to.
    """

    doc_type = _detect_document_type(question, conversation_history)

    if doc_type is None:
        context = (
            "The user asked to generate a document but didn't specify "
            "which kind. NOVA can currently generate: a Purchase Order "
            "document, a Leave Letter, an Expense Report, or Meeting "
            "Minutes."
        )
        confirmation = (
            "I can generate a few kinds of documents for you: a "
            "**Purchase Order** document, a **Leave Letter**, an "
            "**Expense Report**, or **Meeting Minutes**. Which one would "
            "you like?"
        )
        return context, [], None, confirmation

    builder = _BUILDERS[doc_type]

    try:
        result, evidence_or_message, sources, confirmation = builder(question, llm_writer)
    except Exception as error:
        message = (
            f"Document generation failed while building the "
            f"{_DOC_TYPE_LABEL[doc_type]}: {error}"
        )
        return message, [], None, message

    if result is None:
        # evidence_or_message is a plain "nothing found" message here,
        # and confirmation is None - fall back to the same text so the
        # user still gets a clear, non-empty reply.
        return evidence_or_message, [], None, evidence_or_message

    (path, filename) = result
    evidence_lines = evidence_or_message

    context = (
        f"A {_DOC_TYPE_LABEL[doc_type]} document was generated and is "
        f"available for download as '{filename}'. It was built from the "
        f"following real, verified data:\n\n" + "\n".join(evidence_lines)
    )

    file_info = {"path": path, "filename": filename, "doc_type": doc_type}

    return context, sources, file_info, confirmation